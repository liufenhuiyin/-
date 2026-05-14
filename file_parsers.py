# Layer: Adapter
# File: app/adapters/file_parsers.py
# Responsibility: 实现 FileParserProtocol——从各类文件中提取纯文本内容。
#                 每种文件格式对应一个解析器类，由 FileService 按 can_parse() 路由。
# Input:  文件路径（str）
# Output: 提取的纯文本（str），失败时返回空串
# 禁止: 业务逻辑、导入 UI 库

from __future__ import annotations

import json


# ──────────────────────────────────────────────
# 纯文本 / Markdown
# ──────────────────────────────────────────────

class PlainTextParser:
    """处理 .txt / .md / .rst 等纯文本文件。"""

    _EXTENSIONS = {".txt", ".md", ".markdown", ".rst", ".text"}

    def can_parse(self, filename: str) -> bool:
        return _ext(filename) in self._EXTENSIONS

    async def extract_text(self, file_path: str) -> str:
        try:
            with open(file_path, "r", encoding="utf-8", errors="replace") as f:
                return f.read()
        except OSError:
            return ""


# ──────────────────────────────────────────────
# JSON
# ──────────────────────────────────────────────

class JsonParser:
    """处理 .json 文件，格式化后作为文本注入。"""

    def can_parse(self, filename: str) -> bool:
        return _ext(filename) == ".json"

    async def extract_text(self, file_path: str) -> str:
        try:
            with open(file_path, "r", encoding="utf-8", errors="replace") as f:
                data = json.load(f)
            return json.dumps(data, ensure_ascii=False, indent=2)
        except (OSError, json.JSONDecodeError):
            return ""


# ──────────────────────────────────────────────
# Word (.docx)
# ──────────────────────────────────────────────

class DocxParser:
    """处理 .docx 文件，提取正文段落与表格文本。"""

    def can_parse(self, filename: str) -> bool:
        return _ext(filename) == ".docx"

    async def extract_text(self, file_path: str) -> str:
        try:
            import docx  # python-docx
        except ImportError:
            return "[无法解析 .docx：请安装 python-docx]"

        try:
            doc = docx.Document(file_path)
            parts: list[str] = []

            # 段落
            for para in doc.paragraphs:
                text = para.text.strip()
                if text:
                    parts.append(text)

            # 表格
            for table in doc.tables:
                for row in table.rows:
                    row_cells = [cell.text.strip() for cell in row.cells]
                    row_text = " | ".join(c for c in row_cells if c)
                    if row_text:
                        parts.append(row_text)

            return "\n\n".join(parts)
        except Exception:
            return ""


# ──────────────────────────────────────────────
# PDF
# ──────────────────────────────────────────────

class PdfParser:
    """处理 .pdf 文件，逐页提取文本。"""

    def can_parse(self, filename: str) -> bool:
        return _ext(filename) == ".pdf"

    async def extract_text(self, file_path: str) -> str:
        try:
            import pdfplumber
        except ImportError:
            return "[无法解析 .pdf：请安装 pdfplumber]"

        try:
            pages: list[str] = []
            with pdfplumber.open(file_path) as pdf:
                for page in pdf.pages:
                    text = page.extract_text()
                    if text:
                        pages.append(text.strip())
            return "\n\n".join(pages)
        except Exception:
            return ""


# ──────────────────────────────────────────────
# 工厂：返回默认解析器列表（注册顺序即优先级）
# ──────────────────────────────────────────────

def default_parsers() -> list:
    """
    返回默认文件解析器列表，供依赖注入时使用。

    使用方式（在 main.py）：
        from app.adapters.file_parsers import default_parsers
        from app.core import FileService
        file_service = FileService(parsers=default_parsers())
    """
    return [
        PlainTextParser(),
        JsonParser(),
        DocxParser(),
        PdfParser(),
    ]


# ──────────────────────────────────────────────
# 工具
# ──────────────────────────────────────────────

def _ext(filename: str) -> str:
    """提取小写扩展名，如 'report.PDF' → '.pdf'。"""
    import os
    return os.path.splitext(filename)[1].lower()
